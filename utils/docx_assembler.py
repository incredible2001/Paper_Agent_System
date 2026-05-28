"""
论文 DOCX 组装器
将结构化论文数据组装为 Word 文档。
"""

import os
import logging
from datetime import datetime
from typing import Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)


def _set_run_font(run, font_name_en="Times New Roman", font_name_cn="宋体", size=None, bold=None):
    """
    统一设置 run 的字体（中英文分别设置）。

    Args:
        run: docx 的 run 对象
        font_name_en: 英文字体
        font_name_zh: 中文字体
        size: 字体大小
        bold: 是否加粗
    """
    run.font.name = font_name_en
    # 设置中文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    if size:
        run.font.size = size
    if bold is not None:
        run.bold = bold


def _set_paragraph_font(paragraph, font_name_en="Times New Roman", font_name_cn="宋体", size=Pt(12)):
    """
    统一设置段落的字体。

    Args:
        paragraph: docx 的段落对象
        font_name_en: 英文字体
        font_name_zh: 中文字体
        size: 字体大小
    """
    for run in paragraph.runs:
        _set_run_font(run, font_name_en, font_name_cn, size)


def assemble_docx(
    paper: dict,
    output_path: str,
    template_path: Optional[str] = None,
) -> str:
    """
    将结构化论文数据组装为 .docx 文件。

    Args:
        paper: 结构化论文数据，格式:
            {
                "title": str,
                "abstract": str,
                "keywords": list[str],
                "sections": [
                    {"heading": str, "level": int, "content": str},
                    ...
                ],
                "references": list[str],  # 格式化的引用列表
                "tables": list[dict],      # 可选
                "figures": list[dict],     # 可选
            }
        output_path: 输出文件路径
        template_path: 可选的 docx 模板路径

    Returns:
        str: 输出文件的绝对路径
    """
    if template_path and os.path.exists(template_path):
        doc = Document(template_path)
    else:
        doc = Document()
        _setup_default_styles(doc)

    # 清除已有内容 (如果使用模板)
    if template_path:
        for para in doc.paragraphs:
            for run in para.runs:
                run.clear()

    # 标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(paper.get("title", "Untitled"))
    _set_run_font(title_run, size=Pt(16), bold=True)

    # 空行
    doc.add_paragraph()

    # 摘要
    if paper.get("abstract"):
        abstract_heading = doc.add_paragraph()
        abstract_run = abstract_heading.add_run("Abstract")
        _set_run_font(abstract_run, size=Pt(12), bold=True)

        abstract_para = doc.add_paragraph(paper["abstract"])
        abstract_para.paragraph_format.first_line_indent = Inches(0.5)
        _set_paragraph_font(abstract_para)

    # 关键词
    if paper.get("keywords"):
        kw_para = doc.add_paragraph()
        kw_run = kw_para.add_run("Keywords: ")
        _set_run_font(kw_run, size=Pt(12), bold=True)
        kw_content = kw_para.add_run(", ".join(paper["keywords"]))
        _set_run_font(kw_content, size=Pt(12))

    doc.add_paragraph()  # 空行

    # 各节正文
    sections_raw = paper.get("sections", [])
    if isinstance(sections_raw, dict):
        sections_raw = [sections_raw]
    if not isinstance(sections_raw, list):
        sections_raw = []
    for section in sections_raw:
        if not isinstance(section, dict):
            # 跳过非 dict 类型的 section（如字符串）
            if isinstance(section, str) and section.strip():
                p = doc.add_paragraph(section)
                p.paragraph_format.first_line_indent = Inches(0.5)
                _set_paragraph_font(p)
            continue
        heading = section.get("heading", "")
        level = section.get("level", 1)
        content = section.get("content", "")

        # 添加标题
        if level == 1:
            h = doc.add_heading(heading, level=1)
        elif level == 2:
            h = doc.add_heading(heading, level=2)
        else:
            h = doc.add_heading(heading, level=3)

        # 设置标题字体
        for run in h.runs:
            _set_run_font(run, size=run.font.size, bold=True)

        # 添加正文内容
        if content:
            # 按段落分割
            paragraphs = content.split("\n\n")
            for para_text in paragraphs:
                para_text = para_text.strip()
                if para_text:
                    p = doc.add_paragraph(para_text)
                    p.paragraph_format.first_line_indent = Inches(0.5)
                    p.paragraph_format.space_after = Pt(6)
                    _set_paragraph_font(p)

    # 参考文献
    if paper.get("references"):
        doc.add_paragraph()
        ref_heading = doc.add_heading("References", level=1)
        for run in ref_heading.runs:
            _set_run_font(run, size=run.font.size, bold=True)
        for i, ref in enumerate(paper["references"], 1):
            ref_para = doc.add_paragraph()
            ref_para.paragraph_format.left_indent = Inches(0.5)
            ref_para.paragraph_format.first_line_indent = Inches(-0.5)
            ref_run = ref_para.add_run(f"[{i}] {ref}")
            _set_run_font(ref_run, size=Pt(10))
            ref_para.paragraph_format.space_after = Pt(3)

    # 保存
    os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
    doc.save(output_path)

    abs_path = os.path.abspath(output_path)
    logger.info(f"Document saved: {abs_path}")
    return abs_path


def _setup_default_styles(doc: Document) -> None:
    """设置默认的文档样式（中文宋体，英文 Times New Roman）。"""
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置行间距
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def create_draft_filename(project_name: str, version: int, suffix: str = "") -> str:
    """
    生成草稿文件名。

    Args:
        project_name: 项目名称
        version: 版本号
        suffix: 可选后缀 (如 "reviewed", "final")

    Returns:
        str: 文件名 (不含路径)
    """
    timestamp = datetime.now().strftime("%Y%m%d")
    parts = [project_name, f"v{version}", timestamp]
    if suffix:
        parts.append(suffix)
    return "_".join(parts) + ".docx"
